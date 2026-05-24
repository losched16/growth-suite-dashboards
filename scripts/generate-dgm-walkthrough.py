"""
Generates docs/dgm-test-drive.docx — the walkthrough for Desert Garden
Montessori (and the template you'll reuse for every school).

Edit the CONTENT variable below to change the doc.
Re-run: python scripts/generate-dgm-walkthrough.py
Output: docs/dgm-test-drive.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = Path(__file__).resolve().parent.parent / "docs" / "dgm-test-drive.docx"
OUT_PATH.parent.mkdir(exist_ok=True)

# ──────────────────────────────────────────────────────────────────────
# Style helpers
# ──────────────────────────────────────────────────────────────────────

BRAND = RGBColor(0x1D, 0x4E, 0xD8)   # blue-700
MUTED = RGBColor(0x6B, 0x72, 0x80)   # slate-500
DARK  = RGBColor(0x11, 0x18, 0x27)   # slate-900
CODE_BG_GRAY = RGBColor(0xF3, 0xF4, 0xF6)  # slate-100

def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BRAND
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = DARK
    return p

def add_para(doc, text, italic=False, muted=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(11)
    if muted:
        r.font.color.rgb = MUTED
    return p

def add_rich(doc, parts):
    """parts: list of (text, kind) where kind in {'normal','bold','italic','code','link'}"""
    p = doc.add_paragraph()
    for text, kind in parts:
        r = p.add_run(text)
        r.font.size = Pt(11)
        if kind == 'bold':
            r.bold = True
        elif kind == 'italic':
            r.italic = True
        elif kind == 'code':
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        elif kind == 'link':
            r.font.color.rgb = BRAND
            r.underline = True
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.size = Pt(11)

def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(item)
        r.font.size = Pt(11)

def add_callout(doc, title, body, tone='note'):
    """Render an in-line callout. Word doesn't have native callouts so
    we fake it with a single-cell table + colored shading."""
    color_by_tone = {
        'note':    'EFF6FF',  # blue-50
        'warning': 'FFFBEB',  # amber-50
        'success': 'ECFDF5',  # emerald-50
    }
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.0)
    # Shading via XML hack — python-docx doesn't expose cell.shading.
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_by_tone.get(tone, 'EFF6FF'))
    cell._tc.get_or_add_tcPr().append(shd)

    # Title (bold)
    p_title = cell.paragraphs[0]
    r = p_title.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    # Body
    p_body = cell.add_paragraph()
    rb = p_body.add_run(body)
    rb.font.size = Pt(10.5)

def add_table(doc, header, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10.5)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10.5)

def add_spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph('')

def add_hr(doc):
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D1D5DB')  # slate-300
    pbdr.append(bottom)
    p_pr.append(pbdr)

# ──────────────────────────────────────────────────────────────────────
# Build
# ──────────────────────────────────────────────────────────────────────

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

add_h1(doc, "Growth Suite — Test Drive Guide")
add_para(doc, "Desert Garden Montessori · 2026-27 enrollment season", muted=True, italic=True)
add_spacer(doc, 1)

# ─── 1. Overview ───
add_h2(doc, "1. Overview — what you're about to test")
add_para(doc, "Growth Suite gives you two surfaces:")
add_bullets(doc, [
    "Admin (you): a Payments + Forms hub embedded inside your portal at the Payments menu item. It's the same iframe you're seeing now.",
    "Parent Portal: what families log into to see invoices, fill enrollment forms, manage tuition, etc. Lives at growth-suite-parent-portal.vercel.app.",
])
add_para(doc, "For the test drive you'll:")
add_numbered(doc, [
    "Log in as a test parent and see exactly what families see.",
    "Push the enrollment form to all enrolled families via Growth Suite + your email workflow.",
    "Run a couple of test invoices through Stripe (no real money moves — see §4).",
])
add_spacer(doc)
add_hr(doc)

# ─── 2. Test parent ───
add_h2(doc, "2. Logging in as a test parent")
add_para(doc, "Two accounts are ready for you.")

add_h3(doc, "Option A — \"Demo Family\" (recommended)")
add_para(doc, "A fully fake family with one fake student. Anything you submit while logged in as the demo parent shows up in your submissions inbox tagged \"Demo Family\" so it's trivial to filter out / clear later. No real family record is touched.")
add_table(doc, header=["Field", "Value"], rows=[
    ["URL",      "https://growth-suite-parent-portal.vercel.app/login"],
    ["Email",    "demo+desert-garden-montessori@growthsuite.test"],
    ["Password", "demo-2026"],
])
add_spacer(doc)

add_h3(doc, "Option B — Michelle (a real DGM parent)")
add_para(doc, "If you want to test against realistic data (a real student record, a real tuition history), use the seeded credentials for an actual DGM parent. Anything you submit will be a real submission against her family — so let us know after and we'll clear it.")
add_table(doc, header=["Field", "Value"], rows=[
    ["URL",      "https://growth-suite-parent-portal.vercel.app/login"],
    ["Email",    "michellelynnpt@gmail.com"],
    ["Password", "dgm-demo-2026"],
])
add_spacer(doc)

add_h3(doc, "What you'll see after logging in")
add_numbered(doc, [
    "Home — greeting, list of students, list of parents on file.",
    "Pending enrollment forms banner (top of Home) — every active form the family hasn't submitted yet, with one-click links per student.",
    "Forms — full list of forms with per-student completion status.",
    "Billing — invoices owed, tuition plan progress.",
    "Family / Messages / Attendance — additional family-facing surfaces.",
])
add_spacer(doc)

add_h3(doc, "How to create another test parent (for your own future tests)")
add_para(doc, "Self-service \"Spin up a test parent\" inside the admin UI is shipping next week. Until then, your Growth Suite contact can create one for any school in under a minute.")
add_callout(
    doc,
    "Internal note (for our team)",
    "The script lives at scripts/seed-demo-family.mjs. Run with --location <ghl_location_id> to target any school. Creates a Family/Parent/Student trio with email demo+<schoolslug>@growthsuite.test and password demo-2026. Idempotent — re-running just resets the password.",
    tone='note',
)
add_spacer(doc)
add_hr(doc)

# ─── 3. Pushing forms ───
add_h2(doc, "3. Pushing the enrollment form to all enrolled families")
add_para(doc, "The actual email blast happens in your contact workflow (the marketing/CRM layer you already have for newsletters). Growth Suite hands you the URL + email template; the workflow does the fan-out.")

add_h3(doc, "Step-by-step")
add_numbered(doc, [
    "Open the Payments hub → Forms tab. You'll see a 5-step \"How to push a form to families\" guide at the top, plus your forms list.",
    "Verify the form is Active. On the row of the enrollment form, click Edit to open the form editor and toggle Active if needed.",
    "Click \"Send to enrolled families\" on the form row. A panel expands with four sub-steps:",
])
add_bullets(doc, [
    "1️⃣ Link — copy the parent-portal home URL.",
    "2️⃣ Email subject + body — pre-filled with {{contact.first_name}} merge syntax. Two copy buttons.",
    "3️⃣ SMS template — short version for follow-up reminders.",
    "4️⃣ Preview / test — buttons to open the form as the demo parent to verify it looks right.",
])
add_numbered(doc, [
    "In your contact workflow (the email/CRM side): open (or create) the smart list of currently-enrolled families. Create a workflow with one email step. Paste the subject + body you copied. Set the workflow to trigger on smart-list members and run it.",
    "Track completion — back in the Forms tab, the \"N submissions\" chip on each form row opens the submissions inbox.",
])

add_h3(doc, "The submissions inbox")
add_bullets(doc, [
    "Submitted (N) — every family that has filled the form. Click a row to expand the answers.",
    "Not yet submitted (N) — families/students who still need to fill it, with primary parent email + phone copy-paste ready.",
    "A progress bar shows your completion percentage.",
])

add_h3(doc, "How the parent sees it")
add_numbered(doc, [
    "They get the email from your workflow → click the link → arrive at the parent portal /home.",
    "They log in (existing password if returning, or set one if new).",
    "On /home they see an amber \"Pending enrollment forms\" banner at the top showing every form they need to complete, with per-student labels.",
    "Click → fill the form → submit. Done.",
])
add_spacer(doc)
add_hr(doc)

# ─── 4. Stripe ───
add_h2(doc, "4. Stripe Connect — testing now, going live later")

add_h3(doc, "What's set up today")
add_para(doc, "Growth Suite's Stripe Connect platform is in test mode (sandbox). When Desert Garden onboards by clicking \"Connect Stripe\" inside the admin, here's exactly what happens:")
add_numbered(doc, [
    "DGM gets redirected to connect.stripe.com and logs in with their existing Stripe credentials.",
    "Stripe creates a separate test connected account linked to our test platform. It is NOT their live merchant account — it's a parallel test record that lives behind the \"Test mode\" toggle in their Stripe dashboard.",
    "Their real (live) Stripe data is completely untouched. They can keep running their existing business on their live account without overlap.",
])
add_callout(
    doc,
    "Communicating this to DGM",
    "\"When you click Connect Stripe in our system, you'll authorize with your normal Stripe credentials, but the connection lives behind the Test Mode toggle in your Stripe dashboard. None of your real customers, invoices, or balances are affected. You'll see the test connection only if you flip your dashboard to Test Mode.\"",
    tone='note',
)

add_h3(doc, "Q1: Will DGM need to reconnect when we go live?")
add_para(doc, "Yes. Test-mode connected accounts and live-mode connected accounts are completely separate records on Stripe's side. When the platform moves to live mode, DGM will click Connect Stripe one more time and authorize against the live platform — a 30-second OAuth flow.")
add_para(doc, "The good news: their bank account, business info, identity verification all live on DGM's Stripe account (which they own). The \"reconnect\" is just re-establishing the platform-merchant relationship. They won't be re-onboarding from scratch.")

add_h3(doc, "Q2: Could we just connect the live account and run test transactions instead?")
add_para(doc, "No — three reasons:")
add_bullets(doc, [
    "Stripe's test card numbers (4242 4242 4242 4242, etc.) ONLY work in test mode. In live mode the API rejects them with card_declined.",
    "To get test-like behavior in live mode you'd have to use real cards and real (small) amounts. That actually moves money through interchange fees and is awkward to demo.",
    "Connecting the live account now means every test transaction looks real to bookkeeping, the bank, and IRS reporting — and the account picks up real risk-monitoring flags from Stripe.",
])
add_para(doc, "The standard Stripe Connect workflow is always: dev → test → live. We're doing it the right way.")

add_h3(doc, "What testing with the current sandbox looks like")
add_table(doc, header=["Card / rail", "Number", "Behavior"], rows=[
    ["Visa (no auth)",           "4242 4242 4242 4242", "Success"],
    ["Visa (3DS required)",      "4000 0025 0000 3155", "Asks for 3DS challenge"],
    ["Insufficient funds",       "4000 0000 0000 9995", "Declined"],
    ["ACH (routing / account)",  "110000000 / 000123456789", "Success after 4-5 days"],
])
add_para(doc, "Expiry: any future date. CVC: any 3 digits. ZIP: any 5 digits. Test charges show up in the sandbox Stripe dashboard immediately. No real money moves.", muted=True, italic=True)

add_h3(doc, "When DGM is ready to go live")
add_numbered(doc, [
    "Growth Suite's platform moves from test to live (Stripe activation, ~24-48 hours).",
    "DGM clicks Connect Stripe one more time — re-authorizes against the live platform.",
    "That's it. Existing test-mode data (tuition plans, invoices, families) persists in Growth Suite — only the Stripe connection swaps over. New invoices go live, get charged for real.",
])
add_spacer(doc)
add_hr(doc)

# ─── 5. FAQ ───
add_h2(doc, "5. FAQ we expect from DGM")

faq = [
    ("What if our test data mixes with our real data?",
     "It can't. Test-mode and live-mode are two parallel universes in Stripe. The connected-account record, the customers, the invoices, the events — all separate buckets. You see only one universe at a time via the Test Mode toggle in your dashboard."),
    ("Can we keep using our existing Stripe account when we go live?",
     "Yes — that's the whole point of Stripe Connect Standard. DGM owns the Stripe account; Growth Suite just connects to it."),
    ("What about our payouts? Will tuition still hit our bank?",
     "Yes. In live mode, money from invoices goes directly to DGM's Stripe balance and pays out to DGM's bank on DGM's normal schedule. Growth Suite never sees or holds the money — it just orchestrates the invoicing and reads the events."),
    ("What's the platform fee?",
     "$25 one-time family setup fee per new family, plus standard Stripe processing fees (2.9% + $0.30 card / 0.8% ACH). The setup fee toggle is on by default for new invoices; uncheck it per-invoice if a family has already paid it."),
    ("What happens to a family's saved payment methods when we move from test to live?",
     "Saved payment methods are scoped to the test connection. When DGM moves to live, families re-enter their card / bank on the first live invoice (which they save again with the \"Save for future autopay\" box). The autopay schedules persist; only the saved card needs re-entering once."),
]
for q, a in faq:
    p_q = doc.add_paragraph()
    rq = p_q.add_run(q)
    rq.bold = True
    rq.font.size = Pt(11)
    add_para(doc, a)
    add_spacer(doc)

add_hr(doc)

# ─── 6. Quick reference ───
add_h2(doc, "6. Quick reference")
add_table(doc, header=["What", "URL / value"], rows=[
    ["Admin Payments hub",                "Payments menu item inside your portal"],
    ["Parent portal — login",             "https://growth-suite-parent-portal.vercel.app/login"],
    ["Parent portal — home",              "https://growth-suite-parent-portal.vercel.app/home"],
    ["Demo Family email",                 "demo+desert-garden-montessori@growthsuite.test"],
    ["Demo Family password",              "demo-2026"],
    ["Michelle (real parent) email",      "michellelynnpt@gmail.com"],
    ["Michelle (real parent) password",   "dgm-demo-2026"],
    ["Stripe test card",                  "4242 4242 4242 4242 / any future expiry / any 3-digit CVC"],
])

add_spacer(doc, 1)
add_para(doc, "Generated by scripts/generate-dgm-walkthrough.py. Re-run after edits.", muted=True, italic=True)

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
print(f"Size: {OUT_PATH.stat().st_size:,} bytes")
